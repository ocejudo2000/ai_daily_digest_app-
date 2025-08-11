
# ai_daily_digest_app_patched.py
# Streamlit app: AI Daily Digest — patched for Streamlit Cloud caching (returns JSON-serializable objects).
# Includes: custom sources (persisted JSON), executive summary, LinkedIn draft (≤500w) w/ image,
# exports (HTML/PDF/DOCX), email sending, and headless mode for automation.

import os, io, json, time, smtplib, ssl, base64, textwrap, datetime, re, argparse
from email.message import EmailMessage
from email.utils import parsedate_to_datetime
from urllib.parse import urlparse

import requests
import feedparser
from bs4 import BeautifulSoup

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader

# --------------------------- Config ---------------------------
st.set_page_config(page_title="AI Daily Digest", layout="wide")
APP_TITLE = "📰 AI Daily Digest — News & Apps"
SOURCES_FILE = "sources.json"

DEFAULT_SOURCES = [
    # Developer & research blogs
    "https://openai.com/blog/rss.xml",
    "https://www.anthropic.com/news.xml",
    "https://ai.googleblog.com/feeds/posts/default",
    "https://deepmind.google/discover/blog/feed.xml",
    "https://www.microsoft.com/en-us/research/feed/",
    "https://meta.ai/blog/rss.xml",
    "https://stability.ai/blog/rss.xml",
    "https://huggingface.co/blog.xml",
    "https://developer.nvidia.com/blog/feed/",
    "https://machinelearning.apple.com/feed.xml",
    # Product & app aggregators
    "https://www.producthunt.com/topics/artificial-intelligence.rss",
    # News
    "https://www.theverge.com/rss/ai-artificial-intelligence/index.xml",
    "https://www.techcrunch.com/tag/artificial-intelligence/feed/",
    "https://venturebeat.com/category/ai/feed/",
    "https://www.zdnet.com/topic/artificial-intelligence/rss.xml",
    # Papers/abstracts (RSS)
    "https://export.arxiv.org/rss/cs.AI",
    "https://export.arxiv.org/rss/cs.CL",
]

# --------------------------- Helpers ---------------------------

def load_sources():
    try:
        if os.path.exists(SOURCES_FILE):
            with open(SOURCES_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                return list(dict.fromkeys([s for s in data if isinstance(s, str) and s.strip()]))
    except Exception:
        pass
    return DEFAULT_SOURCES.copy()

def save_sources(sources):
    try:
        with open(SOURCES_FILE, "w", encoding="utf-8") as f:
            json.dump(sources, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.warning(f"Could not save sources.json: {e}")
        return False

@st.cache_data(show_spinner=False, ttl=3600)
def fetch_feed(url, timeout=20):
    """Return a JSON-serializable dict to keep cache picklable."""
    try:
        fp = feedparser.parse(url)
        feed_title = str(fp.get("feed", {}).get("title") or url)
        cleaned_entries = []
        for e in fp.get("entries", []):
            entry = {
                "title": str(e.get("title","")),
                "link": str(e.get("link","")),
                "summary": str(e.get("summary","")),
                "published": str(e.get("published","") or e.get("updated","")),
                "published_iso": None,
                "updated_iso": None,
                "has_media": bool(e.get("media_content")),
            }
            try:
                if e.get("published_parsed"):
                    dt = datetime.datetime(*e["published_parsed"][:6], tzinfo=datetime.timezone.utc)
                    entry["published_iso"] = dt.isoformat()
                elif entry["published"]:
                    dt = parsedate_to_datetime(entry["published"])
                    if dt.tzinfo is None:
                        dt = dt.replace(tzinfo=datetime.timezone.utc)
                    entry["published_iso"] = dt.astimezone(datetime.timezone.utc).isoformat()
            except Exception:
                pass
            try:
                if e.get("updated_parsed"):
                    dt = datetime.datetime(*e["updated_parsed"][:6], tzinfo=datetime.timezone.utc)
                    entry["updated_iso"] = dt.isoformat()
            except Exception:
                pass
            cleaned_entries.append(entry)
        return {"feed": {"title": feed_title}, "entries": cleaned_entries}
    except Exception as ex:
        return {"feed": {"title": url}, "entries": [], "error": str(ex)}

def is_recent(entry, days=1):
    """Accept our sanitized entry dict and check recency."""
    try:
        if entry.get("published_iso"):
            dt = datetime.datetime.fromisoformat(entry["published_iso"].replace("Z","+00:00"))
        elif entry.get("updated_iso"):
            dt = datetime.datetime.fromisoformat(entry["updated_iso"].replace("Z","+00:00"))
        elif entry.get("published"):
            dt = parsedate_to_datetime(entry["published"])
        else:
            return True
        now = datetime.datetime.now(datetime.timezone.utc)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=datetime.timezone.utc)
        return (now - dt).days < days
    except Exception:
        return True

def extract_og_image(url, timeout=10):
    try:
        r = requests.get(url, timeout=timeout, headers={"User-Agent":"Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "lxml")
        og = soup.find("meta", property="og:image")
        if og and og.get("content"):
            return og["content"]
        img = soup.find("img")
        if img and img.get("src"):
            src = img["src"]
            if src.startswith("//"): return "https:" + src
            return src
    except Exception:
        return None

def guess_has_video(link):
    return any(k in link for k in ("youtube.com","youtu.be","vimeo.com"))

def clean_text(s, max_chars=None):
    if not s: return ""
    s = BeautifulSoup(s, "lxml").get_text(separator=" ", strip=True)
    s = re.sub(r"\s+", " ", s).strip()
    if max_chars and len(s) > max_chars:
        return s[:max_chars-1] + "…"
    return s

def collect_items(sources, days_back=1, max_items=120):
    items = []
    for src in sources:
        fp = fetch_feed(src)
        feed_title = fp.get("feed",{}).get("title") or src
        for e in fp.get("entries", []):
            title = (e.get("title") or "").strip()
            link = (e.get("link") or "").strip()
            if not title or not link:
                continue
            if days_back and not is_recent(e, days=days_back):
                continue
            summary = clean_text(e.get("summary",""), 400)
            items.append({
                "source": feed_title,
                "title": title,
                "link": link,
                "summary": summary,
                "published": e.get("published_iso") or e.get("published") or "",
                "has_video": guess_has_video(link) or bool(e.get("has_media")),
            })
    seen = set()
    dedup = []
    for it in items:
        if it["link"] in seen: 
            continue
        seen.add(it["link"])
        dedup.append(it)
        if len(dedup) >= max_items:
            break
    return dedup

def call_openai_summarize(bullets, style="executive", max_words=220, api_key=None):
    if not api_key:
        return None
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        text = "\n".join(f"- {b['title']} — {b['link']}" for b in bullets[:20])
        prompt = f"""You're an expert AI industry analyst. Write an executive summary (<= {max_words} words) of today's key AI developments based on the headlines/links below. 
Tone: concise, neutral, boardroom-ready. No fluff. 
Bulleted priorities are allowed for clarity.
HEADLINES:
{text}"""
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.3,
            max_tokens=600,
        )
        return resp.choices[0].message.content.strip()
    except Exception:
        return None

def call_openai_linkedin(items, max_words=500, api_key=None):
    if not api_key: 
        return None
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        points = "\n".join(f"- {clean_text(i['title'],150)} ({i['link']})" for i in items[:8])
        prompt = f"""Write a LinkedIn-style executive post (<= {max_words} words) summarizing today's most important AI launches and research. 
Audience: senior executives. Tone: authoritative, practical, optimistic, jargon-light. 
Include a tight intro, 3–5 bullet takeaways, and a closing call-to-action to follow for updates. No hashtags in body, one line with 3 hashtags at end.
Use these items:
{points}"""
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.4,
            max_tokens=800,
        )
        return resp.choices[0].message.content.strip()
    except Exception:
        return None

def fallback_exec_summary(items, max_words=220):
    bullets = [f"• {clean_text(i['title'],120)} — {i['source']}" for i in items[:10]]
    txt = "Today in AI:\n" + "\n".join(bullets)
    if len(txt.split()) > max_words:
        return " ".join(txt.split()[:max_words]) + "…"
    return txt

def fallback_linkedin(items, max_words=500):
    bullets = "\n".join(f"- {clean_text(i['title'],120)} ({i['link']})" for i in items[:6])
    body = f"""AI weekly snapshot — What matters:\n\n{bullets}\n\nLeaders: pressure-test your AI roadmap against these shifts and double down on high-ROI pilots.\n\n#AI #GenAI #Strategy"""
    if len(body.split()) > max_words:
        return " ".join(body.split()[:max_words]) + "…"
    return body

def build_html(items, exec_summary, linkedin_html, cover_img_url=None):
    rows = []
    for it in items:
        vid = "🎬 " if it["has_video"] else ""
        rows.append(f"<li>{vid}<a href='{it['link']}'>{it['title']}</a> — <em>{it['source']}</em></li>")
    cover = f"<img src='{cover_img_url}' style='max-width:100%;border-radius:8px'/>" if cover_img_url else ""
    html = f"""<html><head><meta charset="utf-8"><title>AI Daily Digest</title></head>
<body>
<h1>{APP_TITLE}</h1>
{cover}
<h2>Executive Summary</h2>
<p>{exec_summary.replace("\n","<br/>")}</p>
<h2>Top Items</h2>
<ol>
{''.join(rows)}
</ol>
<h2>LinkedIn Draft (≤500w)</h2>
<div style="background:#f7f7f7;padding:12px;border-radius:8px">{linkedin_html.replace("\n","<br/>")}</div>
</body></html>"""
    return html.encode("utf-8")

def build_docx(items, exec_summary, linkedin_text, cover_img_bytes=None):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading(APP_TITLE, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cover_img_bytes:
        from docx.shared import Inches
        doc.add_picture(io.BytesIO(cover_img_bytes), width=Inches(6.0))
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(exec_summary)
    doc.add_heading("Top Items", level=1)
    for it in items:
        p = doc.add_paragraph(style=None)
        p.add_run("• ").bold = True
        p.add_run(f"{it['title']} ").bold = True
        p.add_run(f"({it['source']}) ")
        p.add_run(it['link'])
    doc.add_heading("LinkedIn Draft (≤500w)", level=1)
    doc.add_paragraph(linkedin_text)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

def build_pdf(items, exec_summary, linkedin_text, cover_img_bytes=None):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    elems = [Paragraph(APP_TITLE, styles["Title"]), Spacer(1,12)]
    if cover_img_bytes:
        img = Image(ImageReader(io.BytesIO(cover_img_bytes)))
        img.drawWidth = 480
        img.drawHeight = img.drawHeight * (480.0 / img.drawWidth)
        elems += [img, Spacer(1,12)]
    elems += [Paragraph("Executive Summary", styles["Heading1"]),
              Paragraph(exec_summary.replace("\n","<br/>"), styles["BodyText"]),
              Spacer(1,12),
              Paragraph("Top Items", styles["Heading1"])]
    data = [["Title", "Source", "Link"]]
    for it in items[:40]:
        data.append([it["title"], it["source"], it["link"]])
    tbl = Table(data, colWidths=[180,100,220])
    tbl.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.lightgrey),
                             ("BOX",(0,0),(-1,-1),1,colors.black),
                             ("INNERGRID",(0,0),(-1,-1),0.25,colors.grey)]))
    elems += [tbl, Spacer(1,12), Paragraph("LinkedIn Draft (≤500w)", styles["Heading1"]),
              Paragraph(linkedin_text.replace("\n","<br/>"), styles["BodyText"])]
    doc.build(elems)
    buf.seek(0)
    return buf.getvalue()

def download_image(url):
    try:
        r = requests.get(url, timeout=10)
        if r.status_code == 200 and r.content:
            return r.content
    except Exception:
        pass
    return None

def send_email(subject, html_body, attachments, smtp_conf):
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp_conf["from_email"]
    msg["To"] = smtp_conf["to_email"]
    msg.set_content("HTML version attached.")
    msg.add_alternative(html_body.decode("utf-8") if isinstance(html_body, (bytes,bytearray)) else html_body, subtype="html")

    for filename, content, mime in attachments:
        maintype, subtype = mime.split("/",1)
        msg.add_attachment(content, maintype=maintype, subtype=subtype, filename=filename)

    use_tls = smtp_conf.get("use_tls", True)
    context = ssl.create_default_context()
    if use_tls:
        with smtplib.SMTP(smtp_conf["smtp_host"], smtp_conf["smtp_port"]) as server:
            server.starttls(context=context)
            server.login(smtp_conf["smtp_user"], smtp_conf["smtp_pass"])
            server.send_message(msg)
    else:
        with smtplib.SMTP_SSL(smtp_conf["smtp_host"], smtp_conf["smtp_port"], context=context) as server:
            server.login(smtp_conf["smtp_user"], smtp_conf["smtp_pass"])
            server.send_message(msg)

# --------------------------- UI ---------------------------

st.title(APP_TITLE)

with st.sidebar:
    st.header("Filters & Options")
    period = st.selectbox("Time window", ["24 hours","3 days","7 days","14 days"], index=0)
    days_map = {"24 hours":1, "3 days":3, "7 days":7, "14 days":14}
    days_back = days_map[period]
    max_items = st.slider("Max items", 20, 200, 80, 10)
    include_formats = st.multiselect("Export formats", ["HTML","PDF","Word (DOCX)"], default=["HTML","PDF","Word (DOCX)"])

    st.header("Email")
    to_email = st.text_input("Recipient (To)", value=st.secrets.get("email",{}).get("to_email",""))
    st.caption("Configure SMTP in secrets. You can override 'To' here.")

    st.header("Summaries")
    use_openai = st.toggle("Use OpenAI for summaries (if API key in secrets)", value=True)

    st.header("Sources")
    st.caption("Add/remove feeds. Saved to sources.json (local).")
    sources = load_sources()
    show_sources = st.expander("Manage sources", expanded=False)
    with show_sources:
        src_to_add = st.text_input("Add new source (RSS or URL)")
        cols = st.columns([1,1])
        if cols[0].button("➕ Add source", use_container_width=True) and src_to_add:
            if src_to_add not in sources:
                sources.append(src_to_add)
                if save_sources(sources):
                    st.success("Source added.")
        if cols[1].button("↺ Reset to defaults", use_container_width=True):
            sources = DEFAULT_SOURCES.copy()
            save_sources(sources)
            st.success("Sources reset.")
        del_src = st.multiselect("Remove selected", sources, default=[])
        if st.button("🗑️ Remove selected") and del_src:
            sources = [s for s in sources if s not in del_src]
            save_sources(sources)
            st.success("Removed.")

# Fetch
with st.spinner("Fetching feeds…"):
    items = collect_items(sources, days_back=days_back, max_items=max_items)

st.subheader(f"Top {len(items)} items")
for it in items[:10]:
    st.markdown(f"- **[{it['title']}]({it['link']})** — *{it['source']}* {'🎬' if it['has_video'] else ''}")

# Summaries
api_key = st.secrets.get("openai",{}).get("api_key") if use_openai else None
exec_summary = call_openai_summarize(items, api_key=api_key) or fallback_exec_summary(items)
linkedin_text = call_openai_linkedin(items, api_key=api_key) or fallback_linkedin(items)

st.markdown("### Executive Summary")
st.write(exec_summary)

# Cover image
cover_url = extract_og_image(items[0]["link"]) if items else None
cover_bytes = download_image(cover_url) if cover_url else None

st.markdown("### LinkedIn Draft (≤500 words)")
if cover_url:
    st.image(cover_url, caption="Suggested cover image")
st.code(linkedin_text, language="markdown")

# Exports
html_bytes = build_html(items, exec_summary, linkedin_text, cover_img_url=cover_url)
pdf_bytes = build_pdf(items, exec_summary, linkedin_text, cover_img_bytes=cover_bytes)
docx_bytes = build_docx(items, exec_summary, linkedin_text, cover_img_bytes=cover_bytes)

colA, colB, colC = st.columns(3)
with colA:
    st.download_button("⬇️ Download HTML", data=html_bytes, file_name="ai_daily_digest.html", mime="text/html")
with colB:
    st.download_button("⬇️ Download PDF", data=pdf_bytes, file_name="ai_daily_digest.pdf", mime="application/pdf")
with colC:
    st.download_button("⬇️ Download DOCX", data=docx_bytes, file_name="ai_daily_digest.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Email send
st.markdown("---")
st.subheader("Send by Email")
if st.button("✉️ Send Now"):
    smtp = st.secrets.get("email", {})
    if not smtp:
        st.error("Missing [email] settings in secrets.")
    else:
        smtp_conf = {
            "smtp_host": smtp.get("smtp_host"),
            "smtp_port": int(smtp.get("smtp_port",587)),
            "smtp_user": smtp.get("smtp_user"),
            "smtp_pass": smtp.get("smtp_pass"),
            "from_email": smtp.get("from_email"),
            "to_email": to_email or smtp.get("to_email"),
            "use_tls": bool(smtp.get("use_tls", True)),
        }
        if not all([smtp_conf["smtp_host"], smtp_conf["smtp_port"], smtp_conf["smtp_user"], smtp_conf["smtp_pass"], smtp_conf["from_email"], smtp_conf["to_email"]]):
            st.error("SMTP fields missing. Check secrets.")
        else:
            attachments = []
            if "HTML" in include_formats:
                attachments.append(("ai_daily_digest.html", html_bytes, "text/html"))
            if "PDF" in include_formats:
                attachments.append(("ai_daily_digest.pdf", pdf_bytes, "application/pdf"))
            if "Word (DOCX)" in include_formats:
                attachments.append(("ai_daily_digest.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
            subject = f"AI Daily Digest — {datetime.date.today().isoformat()}"
            try:
                send_email(subject, html_bytes, attachments, smtp_conf)
                st.success(f"Email sent to {smtp_conf['to_email']}")
            except Exception as e:
                st.error(f"Failed to send email: {e}")

# Guidance sections
with st.expander("Automation (GitHub Actions) — run daily"):
    st.markdown("""
To automate daily emails, add this workflow in your GitHub repo (replace secrets).

**.github/workflows/daily.yml**
```yaml
name: AI Daily Digest
on:
  schedule:
    - cron: "0 12 * * *"   # every day at 12:00 UTC
  workflow_dispatch:

jobs:
  run:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with:
          python-version: "3.11"
      - run: pip install -r requirements.txt
      - name: Run headless digest
        env:
          EMAIL_SMTP_HOST: ${{ secrets.EMAIL_SMTP_HOST }}
          EMAIL_SMTP_PORT: ${{ secrets.EMAIL_SMTP_PORT }}
          EMAIL_SMTP_USER: ${{ secrets.EMAIL_SMTP_USER }}
          EMAIL_SMTP_PASS: ${{ secrets.EMAIL_SMTP_PASS }}
          EMAIL_FROM: ${{ secrets.EMAIL_FROM }}
          EMAIL_TO: ${{ secrets.EMAIL_TO }}
          OPENAI_API_KEY: ${{ secrets.OPENAI_API_KEY }}
          DAYS_BACK: "1"
          MAX_ITEMS: "80"
        run: python ai_daily_digest_app_patched.py --headless
```

This script already supports the `--headless` flag.
""")

with st.expander("Storage (custom sources)"):
    st.markdown("""
**Local JSON:** This app saves sources to `sources.json`. On Streamlit Cloud, the file system may reset on redeploys.  
**Recommended:** store your sources in a DB (e.g., Supabase, Google Sheets). 
For Sheets, read them with `gspread` using a service account via `st.secrets`.
""")

# --------------------------- Headless mode ---------------------------
def run_headless():
    days_env = int(os.getenv("DAYS_BACK", "1"))
    max_items_env = int(os.getenv("MAX_ITEMS", "80"))
    key = os.getenv("OPENAI_API_KEY", None)
    smtp = dict(
        smtp_host=os.getenv("EMAIL_SMTP_HOST"),
        smtp_port=int(os.getenv("EMAIL_SMTP_PORT","587")),
        smtp_user=os.getenv("EMAIL_SMTP_USER"),
        smtp_pass=os.getenv("EMAIL_SMTP_PASS"),
        from_email=os.getenv("EMAIL_FROM"),
        to_email=os.getenv("EMAIL_TO"),
        use_tls=True,
    )
    sources = load_sources()
    items = collect_items(sources, days_back=days_env, max_items=max_items_env)
    exec_s = call_openai_summarize(items, api_key=key) or fallback_exec_summary(items)
    linkd = call_openai_linkedin(items, api_key=key) or fallback_linkedin(items)
    cover_url = extract_og_image(items[0]["link"]) if items else None
    cover_bytes = download_image(cover_url) if cover_url else None
    html_b = build_html(items, exec_s, linkd, cover_img_url=cover_url)
    pdf_b = build_pdf(items, exec_s, linkd, cover_img_bytes=cover_bytes)
    docx_b = build_docx(items, exec_s, linkd, cover_img_bytes=cover_bytes)
    attachments = [("ai_daily_digest.html", html_b, "text/html"),
                   ("ai_daily_digest.pdf", pdf_b, "application/pdf"),
                   ("ai_daily_digest.docx", docx_b, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")]
    send_email(f"AI Daily Digest — {datetime.date.today().isoformat()}", html_b, attachments, smtp)
    print("Headless digest sent.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--headless", action="store_true", help="Run and email the digest without UI")
    args = parser.parse_args()
    if args.headless:
        run_headless()
